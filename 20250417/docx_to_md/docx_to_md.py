# This program is convert docx files to md files
# Monitoring the docx files in the docx directory
# and convert them to md files in the md directory dynamicly

import docx
import sys, os
import glob
import hashlib

# The main program

scanned_files = {}

# Loop through all the docx files in the docx directory
while True:
    for file in glob.glob("docx/*.docx"):
        if "~$" in file:
            continue

        # Open the docx file
        doc = docx.Document(file)

        # Open the md file
        md_file_path = "md/" + os.path.basename(file).replace(".docx", ".md")

        doc_file_hash = hashlib.sha256('\n'.join([paragraph.text for paragraph in doc.paragraphs]).encode('utf-8')).hexdigest()

        if scanned_files.get(file) == doc_file_hash:
            continue

        # Update scanned_files
        scanned_files[file] = doc_file_hash

        with open(md_file_path, "w") as f:
            # Loop through all the paragraphs in the docx file
            for paragraph in doc.paragraphs:
                # Write the paragraph to the md file
                f.write(paragraph.text + "\n")

        print("Done")
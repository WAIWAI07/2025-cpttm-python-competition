# This program is converts a md file to a docx file

import docx
import sys, os
import glob

# The main program

if __name__ == "__main__":
    files = glob.glob("md/*.md")

    for file in files:
        doc = docx.Document()
        with open(file, "r") as f:
            for line in f.readlines():
                # Edit the line according to the md format
                if line.strip() == "":
                    continue
                if line.startswith("# "):
                    doc.add_heading(line[2:], 0)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], 1)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], 2)
                elif line.startswith("- "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                else:
                    doc.add_paragraph(line)
        doc.save("docx/" + os.path.basename(file).replace(".md", ".docx"))

    print("Done")